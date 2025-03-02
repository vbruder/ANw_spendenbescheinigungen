from concurrent.futures import ThreadPoolExecutor
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from thefuzz import fuzz
import os
from datetime import datetime
from dateutil import parser
from typing import Optional, Dict, List
import msoffcrypto
import io
import openpyxl
import json
import os.path
from docx2pdf import convert
# from tqdm import tqdm
import csv

class DonationReceiptApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Donation Receipt Generator")
        self.root.geometry("1400x800")
        try:
            self.root.iconbitmap('icon.ico')
        except:
            print("Could not find icon.ico in executable path.")

        # Data storage
        self.address_df: Optional[pd.DataFrame] = None
        self.bank_df: Optional[pd.DataFrame] = None
        self.matched_data: List[Dict] = []

        # Config file path
        self.config_file = os.path.join(
            os.path.expanduser("."), ".donation_receipt_config.json"
        )

        # Load saved paths
        self.load_config()

        self.setup_ui()
        self.root.after(100, self.set_initial_focus)

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def set_initial_focus(self):
        """Set initial focus on the load button"""
        if hasattr(self, "load_button"):
            self.load_button.focus_set()

    def load_config(self):
        """Load saved file paths from config file"""
        self.config = {
            "address_file": "",
            "bank_file": "",
            "template_file": "",
            "output_dir": "",
            "log_dir": "",
            "output_dir_pdf": "",
            "geometry": "",
        }

        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, "r", encoding="utf-8") as f:
                    saved_config = json.load(f)
                    self.config.update(saved_config)
                    self.root.geometry(self.config["geometry"])
        except Exception as e:
            print(f"Error loading config: {str(e)}")

    def save_config(self):
        """Save current file paths to config file"""
        try:
            config_to_save = {
                "address_file": self.address_file_var.get(),
                "bank_file": self.bank_file_var.get(),
                "template_file": self.template_file_var.get(),
                "output_dir": self.output_dir_var.get(),
                "log_dir": self.log_dir_var.get(),
                "output_dir_pdf": self.output_dir_pdf_var.get(),
                "geometry": self.root.geometry(),
            }

            with open(self.config_file, "w", encoding="utf-8") as f:
                json.dump(config_to_save, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving config: {str(e)}")

    def setup_ui(self):
        """Create the main UI components"""
        # Create main container with padding
        main_container = ttk.Frame(self.root, padding="10")
        main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # File Selection Frame
        self.create_file_selection_frame(main_container)

        # Data View Frame
        self.create_data_view_frame(main_container)

        # Output Options Frame
        self.create_output_options_frame(main_container)

        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_container.columnconfigure(1, weight=1)
        main_container.rowconfigure(1, weight=1)

    def create_file_selection_frame(self, parent):
        """Create the file selection section"""
        file_frame = ttk.LabelFrame(parent, text="File Selection", padding="5")
        file_frame.grid(
            row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10)
        )

        # Address File Selection
        ttk.Label(file_frame, text="Address File (.xlsx):").grid(
            row=0, column=0, sticky=tk.W
        )
        self.address_file_var = tk.StringVar(value=self.config["address_file"])
        ttk.Entry(file_frame, textvariable=self.address_file_var, width=150).grid(
            row=0, column=1, padx=5
        )
        ttk.Button(file_frame, text="Browse", command=self.browse_address_file).grid(
            row=0, column=2
        )

        # Password Field
        ttk.Label(file_frame, text="Password:").grid(row=0, column=3, padx=(20, 5))
        self.password_var = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.password_var, show="*").grid(
            row=0, column=4, padx=5
        )

        # Bank File Selection
        ttk.Label(file_frame, text="Bank Statement (.csv):").grid(
            row=1, column=0, sticky=tk.W, pady=(5, 0)
        )
        self.bank_file_var = tk.StringVar(value=self.config["bank_file"])
        ttk.Entry(file_frame, textvariable=self.bank_file_var, width=150).grid(
            row=1, column=1, pady=(5, 0), padx=5
        )
        ttk.Button(file_frame, text="Browse", command=self.browse_bank_file).grid(
            row=1, column=2, pady=(5, 0)
        )

        # Load Data Button
        self.load_button = ttk.Button(
            file_frame, text="Load Data", command=self.load_data
        )
        self.load_button.grid(row=1, column=3, columnspan=2, pady=(5, 0))

    def create_data_view_frame(self, parent):
        """Create the data view section with the table"""
        data_frame = ttk.LabelFrame(parent, text="Matched Data", padding="5")
        data_frame.grid(
            row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10)
        )

        # Create Treeview
        self.tree = ttk.Treeview(
            data_frame,
            columns=(
                "donor_name",
                "matched_name",
                "street",
                "postal_code",
                "city",
                "amount",
                "date",
                "match_score",
                "purpose",
            ),
            show="headings",
            selectmode="browse",
        )

        # Configure columns with headers
        column_headers = {
            "donor_name": "Name on Bank Statement",
            "matched_name": "Matched Name",
            "street": "Street",
            "postal_code": "Postal Code",
            "city": "City",
            "amount": "Amount (EUR)",
            "date": "Date",
            "match_score": "Match Score",
            "purpose": "Purpose",
        }

        # Set up columns with initial width and headers
        for col, header in column_headers.items():
            self.tree.heading(col, text=header)
            self.tree.column(col, width=100)  # Default width, will be adjusted later

        # Add scrollbars
        y_scroll = ttk.Scrollbar(
            data_frame, orient=tk.VERTICAL, command=self.tree.yview
        )
        x_scroll = ttk.Scrollbar(
            data_frame, orient=tk.HORIZONTAL, command=self.tree.xview
        )
        self.tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)

        # Grid layout
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        y_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        x_scroll.grid(row=1, column=0, sticky=(tk.W, tk.E))

        # Configure grid weights
        data_frame.columnconfigure(0, weight=1)
        data_frame.rowconfigure(0, weight=1)

        # Bind events
        self.tree.bind("<Double-1>", self.edit_entry)
        self.tree.bind("<Delete>", lambda event: self.remove_entry()) 

        # Add buttons for data management
        button_frame = ttk.Frame(data_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(5, 0))

        ttk.Button(button_frame, text="Add New Entry", command=self.add_new_entry).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(
            button_frame, text="Update Address File", command=self.update_address_file
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            button_frame, text="Remove Entry", command=self.remove_entry
        ).pack(side=tk.LEFT, padx=5)

    def remove_entry(self):
        """Remove the selected entry from the table and matched_data"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("Warning", "Please select an entry to remove.")
            return

        if messagebox.askyesno("Confirm Removal", "Are you sure you want to remove this entry?"):
            for item in selected_items:
                # Get the index of the item in the tree
                idx = self.tree.index(item)
                
                # Remove from matched_data list
                if 0 <= idx < len(self.matched_data):
                    self.matched_data.pop(idx)
                
                # Remove from treeview
                self.tree.delete(item)

    def convert_to_pdf(self, docx_path, output_dir):
        """
        Convert a single Word document to PDF.
        
        Args:
            docx_path (str): Path to the Word document
            output_dir (str): Directory where PDF should be saved
        Returns:
            str: Path to the generated PDF file
        """
        try:
            # Create PDF filename from Word filename
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            convert(docx_path, pdf_path)
            
            return pdf_path
        except Exception as e:
            print(f"Error converting {docx_path} to PDF: {str(e)}")
            return None

    def batch_convert_to_pdf(self):
        """
        Convert all Word documents in the output directory to PDF.
        
        Args:
            output_dir (str): Directory containing Word documents
        Returns:
            tuple: (successful_conversions, failed_conversions)
        """
        pdf_dir = self.output_dir_pdf_var.get()
        
        # Get all Word documents in the directory
        docx_files = [f for f in os.listdir(self.output_dir_var.get()) if f.endswith('.docx')]
        
        successful = []
        failed = []
        
        # Create progress dialog
        progress_dialog = PdfConvertProgressDialog(self.root, len(docx_files))
        progress_dialog.title("Converting to PDF")
        
        try:
            # Process files sequentially with progress tracking
            for i, docx_file in enumerate(docx_files):
                docx_path = os.path.join(self.output_dir_var.get(), docx_file)
                
                # Update progress message
                progress_dialog.label.config(text=f"Converting {docx_file}...")
                
                try:
                    pdf_path = self.convert_to_pdf(docx_path, pdf_dir)
                    if pdf_path:
                        successful.append(docx_file)
                    else:
                        failed.append(docx_file)
                except Exception as e:
                    print(f"\nError converting {docx_file}: {str(e)}")
                    failed.append(docx_file)
                
                # Update progress bar
                progress_dialog.update(i + 1)
                
        finally:
            progress_dialog.destroy()
        
        # Show completion message
        total = len(docx_files)
        success_count = len(successful)
        fail_count = len(failed)
        
        message = f"Conversion complete!\n\n" \
                f"Successfully converted: {success_count}/{total}\n" \
                f"Failed conversions: {fail_count}/{total}"
        
        if failed:
            message += "\n\nFailed files:\n" + "\n".join(failed)
        
        messagebox.showinfo("Conversion Complete", message)
        
        return successful, failed

    def create_output_options_frame(self, parent):
        """Create the output options section"""
        output_frame = ttk.LabelFrame(parent, text="Output Options", padding="5")
        output_frame.grid(
            row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10)
        )

        # Template File Selection
        ttk.Label(output_frame, text="Template File (.docx)").grid(
            row=0, column=0, sticky=tk.W
        )
        self.template_file_var = tk.StringVar(value=self.config["template_file"])
        ttk.Entry(output_frame, textvariable=self.template_file_var, width=150).grid(
            row=0, column=1, padx=5
        )
        ttk.Button(output_frame, text="Browse", command=self.browse_template_file).grid(
            row=0, column=2
        )

        # Output Directory Selection
        ttk.Label(output_frame, text="Output Directory (docx)").grid(
            row=1, column=0, sticky=tk.W, pady=(5, 0)
        )
        self.output_dir_var = tk.StringVar(value=self.config["output_dir"])
        ttk.Entry(output_frame, textvariable=self.output_dir_var, width=150).grid(
            row=1, column=1, padx=5, pady=(5, 0)
        )
        ttk.Button(output_frame, text="Browse", command=self.browse_output_dir).grid(
            row=1, column=2, pady=(5, 0)
        )
        # generate receipts
        ttk.Button(
            output_frame, text="Generate Receipts", command=self.generate_receipts
        ).grid(row=1, column=3,  pady=(5, 0))

        # Log Output Directory Selection
        ttk.Label(output_frame, text="Log Output Directory").grid(row=2, column=0, sticky=tk.W, pady=(5, 0))
        self.log_dir_var = tk.StringVar(value=self.config["log_dir"])
        ttk.Entry(output_frame, textvariable=self.log_dir_var, width=150).grid(row=2, column=1, padx=5, pady=(5, 0))
        ttk.Button(output_frame, text="Browse", command=self.browse_log_dir).grid(row=2, column=2, pady=(5, 0))

        # Output Directory PDF
        ttk.Label(output_frame, text="Output Directory (pdf)").grid(
            row=3, column=0, sticky=tk.W, pady=(5, 0)
        )
        self.output_dir_pdf_var = tk.StringVar(value=self.config["output_dir_pdf"])
        ttk.Entry(output_frame, textvariable=self.output_dir_pdf_var, width=150).grid(
            row=3, column=1, padx=5, pady=(5, 0)
        )
        ttk.Button(output_frame, text="Browse", command=self.browse_output_dir_pdf).grid(
            row=3, column=2, pady=(5, 0)
        )
        ttk.Button(
            output_frame, text="Convert to PDFs", command=self.convert_to_pdfs
        ).grid(row=3, column=3, pady=(5, 0))

    def browse_template_file(self):
        """Open file dialog for template file selection"""
        filename = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialdir=(
                os.path.dirname(self.template_file_var.get())
                if self.template_file_var.get()
                else None
            ),
        )
        if filename:
            self.template_file_var.set(filename)
            self.save_config()

    def browse_address_file(self):
        """Open file dialog for address file selection"""
        filename = filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            initialdir=(
                os.path.dirname(self.address_file_var.get())
                if self.address_file_var.get()
                else None
            ),
        )
        if filename:
            self.address_file_var.set(filename)
            self.save_config()

    def browse_bank_file(self):
        """Open file dialog for bank file selection"""
        filename = filedialog.askopenfilename(
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            initialdir=(
                os.path.dirname(self.bank_file_var.get())
                if self.bank_file_var.get()
                else None
            ),
        )
        if filename:
            self.bank_file_var.set(filename)
            self.save_config()

    def browse_output_dir(self):
        """Open directory dialog for output selection"""
        directory = filedialog.askdirectory(
            initialdir=self.output_dir_var.get() if self.output_dir_var.get() else None
        )
        if directory:
            self.output_dir_var.set(directory)
            self.save_config()

    def browse_output_dir_pdf(self):
        """Open directory dialog for output pdf selection"""
        directory = filedialog.askdirectory(
            initialdir=self.output_dir_pdf_var.get() if self.output_dir_pdf_var.get() else None
        )
        if directory:
            self.output_dir_pdf_var.set(directory)
            self.save_config()

    def browse_log_dir(self):
        """Open directory dialog for output log selection"""
        directory = filedialog.askdirectory(
            initialdir=self.log_dir_var.get() if self.log_dir_var.get() else None
        )
        if directory:
            self.log_dir_var.set(directory)
            self.save_config()

    def load_data(self):
        """Load and process the data files"""
        try:
            # Create progress dialog
            progress = LoadingProgressDialog(self.root)
            self.root.update()

            # Show loading indicator
            self.root.config(cursor="wait")

            # Load address file
            progress.update_status("Loading address file...", 10)
            self.address_df = self.load_address_data(
                self.address_file_var.get(), self.password_var.get()
            )

            # Load bank file
            progress.update_status("Loading bank statement file...", 30)
            self.bank_df = self.load_bank_data(self.bank_file_var.get())

            # Process matches
            progress.update_status("Processing matches...", 50)
            total_records = len(self.bank_df)

            self.matched_data = []
            for i, donation in enumerate(self.bank_df.iterrows()):
                if donation[1]["Betrag"] <= 0:
                    continue

                donor_name = donation[1]["Beguenstigter/Zahlungspflichtiger"]
                amount = float(str(donation[1]["Betrag"]).replace(",", "."))
                date = self.format_date_str(donation[1]["Buchungstag"])
                purpose = donation[1]["Verwendungszweck"]

                # Find best match
                best_match, score = self.find_best_match(donor_name)

                match_data = {
                    "donor_name": donor_name,
                    "matched_name": (
                        best_match["Name"] if best_match is not None else ""
                    ),
                    "street": best_match["Straße"] if best_match is not None else "",
                    "postal_code": best_match["PLZ"] if best_match is not None else "",
                    "city": best_match["Ort"] if best_match is not None else "",
                    "amount": f"{amount:.2f}",
                    "date": date,
                    "match_score": f"{score:.1f}" if score > 0 else "0.0",
                    "purpose": purpose,
                }

                self.matched_data.append(match_data)

                # Update progress
                progress_value = 50 + (i / total_records * 40)  # Scale from 50 to 90
                progress.update_status(
                    f"Matching records... ({i+1}/{total_records})", progress_value
                )

            # Update table
            progress.update_status("Updating display...", 90)
            self.update_table()

            progress.update_status("Complete!", 100)
            progress.destroy()

        except Exception as e:
            if progress:
                progress.destroy()
            messagebox.showerror("Error", f"Error loading data: {str(e)}")
        finally:
            self.root.config(cursor="")

    def load_address_data(
        self, excel_path: str, password: Optional[str]
    ) -> pd.DataFrame:
        """
        Load the address Excel file, handling both encrypted and unencrypted files.
        First tries to load the file directly, falls back to decryption if needed.
        
        Args:
            excel_path: Path to the Excel file
            password: Optional password for encrypted files
        
        Returns:
            pandas.DataFrame: Loaded data from the Excel file
        """
        try:
            # First try to load the file directly
            try:
                workbook = openpyxl.load_workbook(filename=excel_path)
            except:
                try:
                    # If direct load fails, try decryption
                    decrypted_workbook = io.BytesIO()
                    with open(excel_path, "rb") as file:
                        office_file = msoffcrypto.OfficeFile(file)
                        if password:
                            office_file.load_key(password=password)
                        office_file.decrypt(decrypted_workbook)
                    workbook = openpyxl.load_workbook(filename=decrypted_workbook)
                except:
                    raise Exception(f"Decryption failed, check the entered password.")

            sheet = workbook.active
            data = []
            headers = [cell.value for cell in sheet[1]]

            for row in sheet.iter_rows(min_row=2):
                row_data = {}
                for header, cell in zip(headers, row):
                    row_data[header] = cell.value
                data.append(row_data)

            return pd.DataFrame(data)

        except Exception as e:
            raise Exception(f"Error reading Excel file: {str(e)}")

    def load_bank_data(self, csv_path: str) -> pd.DataFrame:
        """Load the bank CSV file"""
        encodings = ["utf-8", "iso-8859-1", "cp1252", "latin1"]

        for encoding in encodings:
            try:
                with open(csv_path, "r", encoding=encoding) as file:
                    first_line = file.readline()
                    delimiter = ";" if ";" in first_line else ","

                df = pd.read_csv(
                    csv_path,
                    sep=delimiter,
                    encoding=encoding,
                    decimal=",",
                )

                return df[
                    ["Buchungstag", "Beguenstigter/Zahlungspflichtiger", "Betrag", "Verwendungszweck"]
                ]

            except UnicodeDecodeError:
                continue
            except Exception as e:
                continue

        raise ValueError(
            "Could not read the CSV file with any of the attempted encodings"
        )

    def log_receipt(self, receipt_data):
        """Log receipt information to year-specific CSV files in the selected directory."""
        # Determine the year from the donation date
        donation_date = datetime.strptime(receipt_data['donation_date'], "%d.%m.%Y")
        year = donation_date.year
        log_file = os.path.join(self.log_dir_var.get(), f"spendenbescheinigungen_{year}.csv")

        headers = [
            'Date Generated',
            'Donor Name',
            'Street',
            'Postal Code',
            'City',
            'Donation Amount',
            'Donation Amount Words',
            'Donation Date',
            'Receipt Filename'
        ]

        # Check if the file exists; create it with headers if it doesn't
        if not os.path.exists(log_file):
            with open(log_file, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f, delimiter=';')
                writer.writerow(headers)

        # Check for duplicate entries
        try:
            with open(log_file, 'r', encoding='utf-8-sig') as f:
                existing_data = list(csv.reader(f, delimiter=';'))
        except FileNotFoundError:
            existing_data = []

        new_entry = [
            receipt_data['generation_date'],
            receipt_data['donor_name'],
            receipt_data['street'],
            receipt_data['postal_code'],
            receipt_data['city'],
            receipt_data['amount'],
            receipt_data['amount_words'],
            receipt_data['donation_date'],
            receipt_data['filename']
        ]

        # Append entry if not already present
        if new_entry not in existing_data:
            with open(log_file, 'a', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f, delimiter=';')
                writer.writerow(new_entry)


    def process_matches(self):
        """Process and match the loaded data"""
        self.matched_data = []

        for _, donation in self.bank_df.iterrows():
            if donation["Betrag"] <= 0:
                continue

            donor_name = donation["Beguenstigter/Zahlungspflichtiger"]
            amount = float(str(donation["Betrag"]).replace(",", "."))
            date = donation["Buchungstag"]
            purpose = donation["Verwendungszweck"]

            # Find best match
            best_match, score = self.find_best_match(donor_name)

            match_data = {
                "donor_name": donor_name,
                "matched_name": best_match["Name"] if best_match is not None else "",
                "street": best_match["Straße"] if best_match is not None else "",
                "postal_code": best_match["PLZ"] if best_match is not None else "",
                "city": best_match["Ort"] if best_match is not None else "",
                "amount": f"{amount:.2f}",
                "date": date,
                "match_score": f"{score:.1f}" if score > 0 else "0.0",
                "purpose": purpose,
            }

            self.matched_data.append(match_data)

    def split_multiple_names(self, full_name):
        """
        Split a string containing multiple names into separate names.
        Handles various formats and separators.
        """
        full_name = str(full_name).strip()

        # List of possible separators
        separators = [
            " Und ",
            " und ",
            " U. ",
            " u. ",
            " And ",
            " and ",
            " & ",
            " + ",
            "   ",
        ]

        # First try explicit separators
        for sep in separators:
            if sep in full_name:
                split_names = full_name.split(sep)
                if len(split_names[0].split(" ")) == 1:
                    split_names = [
                        split_names[0] + " " + split_names[-1].split(" ")[-1],
                        split_names[1],
                    ]
                return [name.strip() for name in split_names]

        # If no explicit separator, try to detect multiple full names
        # by looking for patterns like multiple last names
        words = full_name.split(" ")
        if len(words) >= 4:  # Minimum 4 words needed for 2 full names
            # Try to find repeated last names
            last_name = words[-1]
            for i in range(len(words) - 2, 0, -1):
                if words[i] == last_name:
                    return [" ".join(words[:i+1]), " ".join(words[i+1:])]

        # If no pattern found, return as single name
        return [full_name]

    def normalize_name(self, name):
        """
        Normalize name for comparison, handling different formats.
        """
        name = str(name).strip()

        # Handle "last_name, first_name" format
        if "," in name:
            parts = name.split(",")
            if len(parts) == 2:
                last_name = parts[0].strip()
                first_name = parts[1].strip()
                return f"{first_name} {last_name}"

        return name

    def find_best_match(self, donor_name, threshold=80):
        """
        Find the best matching address using fuzzy matching.
        Handles multiple names and tries various matching strategies.
        """
        best_score = 0
        best_match = None
        has_best_match = False
        original_donor_name = donor_name
        matched_name = None

        # turn posible all caps into regular title format
        formatted_name = donor_name.title()
        # normalize "last_name, firstname" format
        donor_name_normalized = self.normalize_name(formatted_name)

        # Split into potential multiple names
        donor_names = self.split_multiple_names(donor_name_normalized)
        print("")
        if len(donor_names) > 1:
            print(f"Split '{donor_name_normalized}' into: {donor_names}")

        # Try matching each name individually and combined
        for name in donor_names:
            normalized_name = self.normalize_name(name)
            print(normalized_name)

            for _, row in self.address_df.iterrows():
                list_name_raw = str(row["Name"])
                list_name_normalized_raw = self.normalize_name(list_name_raw)

                list_names = self.split_multiple_names(list_name_raw)
                list_names_normalized = self.split_multiple_names(
                    list_name_normalized_raw
                )

                for [list_name, list_name_normalized] in zip(
                    list_names, list_names_normalized
                ):
                    # Try different matching combinations
                    scores = [
                        fuzz.ratio(name.lower(), list_name.lower()),
                        fuzz.ratio(normalized_name.lower(), list_name.lower()),
                        fuzz.ratio(
                            normalized_name.lower(), list_name_normalized.lower()
                        ),
                        fuzz.token_sort_ratio(name.lower(), list_name.lower()),
                        fuzz.token_sort_ratio(
                            normalized_name.lower(), list_name_normalized.lower()
                        ),
                    ]

                    max_score = max(scores)

                    if max_score > best_score and max_score >= threshold:
                        best_score = max_score
                        best_match = row
                        has_best_match = True
                        matched_name = name

        if len(donor_names) > 1 and not has_best_match:
            # If no match found and we have multiple names,
            # try matching the combined names
            combined_name = " ".join(donor_names)

            for _, row in self.address_df.iterrows():
                list_name = str(row["Name"])
                list_name_normalized = self.normalize_name(list_name)

                scores = [
                    fuzz.ratio(combined_name.lower(), list_name.lower()),
                    fuzz.ratio(combined_name.lower(), list_name_normalized.lower()),
                    fuzz.token_sort_ratio(combined_name.lower(), list_name.lower()),
                    fuzz.token_sort_ratio(
                        combined_name.lower(), list_name_normalized.lower()
                    ),
                ]

                max_score = max(scores)

                if max_score > best_score and max_score >= threshold:
                    best_score = max_score
                    best_match = row
                    matched_name = combined_name

        if best_match is not None:
            print(
                f"Found match for '{matched_name}' in '{best_match['Name']}' with score {best_score}"
            )
            if matched_name != original_donor_name:
                print(f"Note: Matched using partial name from '{original_donor_name}'")
        else:
            print(f"No match found for any name in '{original_donor_name}'")

        return best_match, best_score

    def update_table(self):
        """Update the treeview with matched data"""
        # Clear existing items
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Add new items
        for data in self.matched_data:
            # Determine the appropriate tag based on match score
            if not data["matched_name"]:
                tags = ("unmatched",)
            elif float(data["match_score"]) < 80:
                tags = ("low_score",)
            elif float(data["match_score"]) < 95:
                tags = ("medium_score",)
            else:
                tags = ()

            self.tree.insert(
                "",
                "end",
                values=(
                    data["donor_name"],
                    data["matched_name"],
                    data["street"],
                    data["postal_code"],
                    data["city"],
                    data["amount"],
                    data["date"],
                    data["match_score"],
                    data["purpose"],
                ),
                tags=tags,
            )

        # Configure tag colors
        self.tree.tag_configure("unmatched", background="#ffcccc")  # Light red for unmatched
        self.tree.tag_configure("low_score", background="#ff6666")  # Red for low scores (< 80)
        self.tree.tag_configure("medium_score", background="#ffff99")  # Yellow for medium scores (80-95)


        # Auto-adjust column widths
        self.adjust_column_widths()

    def adjust_column_widths(self, padding=20):
        """Adjust column widths based on content"""
        for column in self.tree["columns"]:
            # Get width of column header
            header = self.tree.heading(column)["text"]
            max_width = self.get_text_width(header)

            # Check content width for each item
            for item in self.tree.get_children():
                cell_value = str(self.tree.set(item, column))
                width = self.get_text_width(cell_value)
                max_width = max(max_width, width)

            # Set column width with padding
            self.tree.column(column, width=max_width + padding)

    def get_text_width(self, text, font_family="TkDefaultFont", font_size=10):
        """Calculate pixel width of text"""
        test_label = tk.Label(self.root, text=text, font=(font_family, font_size))
        width = test_label.winfo_reqwidth()
        test_label.destroy()
        return width

    def edit_entry(self, event):
        """Handle double-click to edit entry"""
        item = self.tree.selection()[0]
        column = self.tree.identify_column(event.x)

        # Get current values
        values = self.tree.item(item)["values"]

        # Create edit dialog
        dialog = EditDialog(self.root, self.address_df, values)
        self.root.wait_window(dialog)

        if dialog.result:
            # Update treeview
            self.tree.item(item, values=dialog.result)

            # Update matched_data
            idx = self.tree.index(item)
            self.matched_data[idx].update(
                {
                    "matched_name": dialog.result[1],
                    "street": dialog.result[2],
                    "postal_code": dialog.result[3],
                    "city": dialog.result[4],
                }
            )

    def add_new_entry(self):
        """Add a new address entry"""
        dialog = EditDialog(self.root, self.address_df)
        self.root.wait_window(dialog)

        if dialog.result:
            # Add to treeview
            self.tree.insert("", "end", values=dialog.result)

            # Add to matched_data
            self.matched_data.append(
                {
                    "donor_name": dialog.result[0],
                    "matched_name": dialog.result[1],
                    "street": dialog.result[2],
                    "postal_code": dialog.result[3],
                    "city": dialog.result[4],
                    "amount": dialog.result[5],
                    "date": dialog.result[6],
                    "match_score": dialog.result[7],
                    "purpose": "",
                }
            )

    def update_address_file(self):
        """Update the address Excel file with new/modified entries"""
        try:
            # Create a backup of the original file
            backup_path = self.address_file_var.get() + ".backup"
            if not os.path.exists(backup_path):
                import shutil

                shutil.copy2(self.address_file_var.get(), backup_path)

            # Get all unique matched names that aren't in the address file
            current_names = set(self.address_df["Name"].astype(str))
            new_entries = []

            for data in self.matched_data:
                if data["matched_name"] and data["matched_name"] not in current_names:
                    new_entries.append(
                        {
                            "Name": data["matched_name"],
                            "Straße": data["street"],
                            "PLZ": data["postal_code"],
                            "Ort": data["city"],
                        }
                    )

            if new_entries:
                # Add new entries to the DataFrame
                new_df = pd.DataFrame(new_entries)
                self.address_df = pd.concat(
                    [self.address_df, new_df], ignore_index=True
                )

                # Save updated DataFrame to Excel
                self.address_df.to_excel(self.address_file_var.get(), index=False)
                messagebox.showinfo(
                    "Success",
                    f"Added {len(new_entries)} new entries to the address file.",
                )
            else:
                messagebox.showinfo(
                    "Info", "No new entries to add to the address file."
                )

        except Exception as e:
            messagebox.showerror("Error", f"Error updating address file: {str(e)}")

    def generate_receipts(self):
        """Generate donation receipts for all matched entries"""
        if not self.matched_data:
            messagebox.showerror("Error", "No data loaded to generate receipts from.")
            return

        template_path = self.template_file_var.get()
        if not template_path:
            messagebox.showerror("Error", "Please select a template file.")
            return

        output_dir = self.output_dir_var.get()
        if not output_dir:
            messagebox.showerror("Error", "Please select an output directory.")
            return

        try:
            # Create output directory if it doesn't exist
            os.makedirs(output_dir, exist_ok=True)

            # Create progress dialog
            progress_dialog = ProgressDialog(self.root, len(self.matched_data))

            for i, data in enumerate(self.matched_data):
                if data["matched_name"]:  # Only generate for matched entries
                    try:
                        self.generate_single_receipt(data, output_dir, template_path)
                    except Exception as e:
                        print(
                            f"Error generating receipt for {data['donor_name']}: {str(e)}"
                        )

                progress_dialog.update(i + 1)

            progress_dialog.destroy()
            messagebox.showinfo("Success", "Receipt generation complete!")

        except Exception as e:
            messagebox.showerror("Error", f"Error generating receipts: {str(e)}")

    def convert_to_pdfs(self):
        """Convert all docs in the docx output directory to PDF files"""
        input_dir = self.output_dir_var.get()
        if not input_dir:
            messagebox.showerror("Error", "Please select an docx output directory where the files to convert are located.")
            return

        output_dir = self.output_dir_pdf_var.get()
        if not output_dir:
            messagebox.showerror("Error", "Please select an output directory for the generated PDF files.")
            return

        try:
            # Create output directory if it doesn't exist
            os.makedirs(output_dir, exist_ok=True)

            # Create progress dialog
            # progress_dialog = ProgressDialog(self.root, len(self.matched_data))

            self.batch_convert_to_pdf()

            #     progress_dialog.update(i + 1)

            # progress_dialog.destroy()
            messagebox.showinfo("Success", "PDF conversion complete!")

        except Exception as e:
            messagebox.showerror("Error", f"Error converting docs to PDFs: {str(e)}")

    def format_date(self, date_int: int):
        """
        Convert date from integer format (DDMMYY) to German date string (DD.MM.YYYY).

        Args:
            date_int (int): Date as integer in format DDMMYY (e.g., 130125 for 13.01.25)
        Returns:
            str: Date in format 'DD.MM.YYYY'
        """
        try:
            # Convert integer to string and pad with leading zeros if necessary
            date_str = str(date_int)
            if len(date_str) < 6:
                date_str = str(date_int).zfill(6)
            elif len(date_str) == 7:
                date_str = str(date_int).zfill(8)

            # Extract components
            day = date_str[:2]
            month = date_str[2:4]
            year = date_str[4:]

            # Convert two-digit year to four-digit year
            year_int = int(year)
            if year_int < 100:
                year = f"20{year}"

            # Return formatted date
            return f"{day}.{month}.{year}"

        except Exception as e:
            print(f"Error formatting date {date_int}: {str(e)}")
            return str(date_int)

    def format_date_str(self, date_str: str):
        """Parse a date string and format it as dd.mm.YYYY."""
        try:
            # Use dateutil's parser to handle different formats
            parsed_date = parser.parse(date_str, dayfirst=True)
            # Format the date as dd.mm.YYYY
            formatted_date = parsed_date.strftime("%d.%m.%Y")
            return formatted_date
        except ValueError:
            return "Invalid date format"

    def generate_single_receipt(self, data, output_dir, template_path):
        """Generate a single donation receipt"""
        from docx import Document
        from datetime import datetime
        import locale

        doc = Document(template_path)

        desired_locales = ["de_DE.UTF-8", "de_DE", "de_de", "German"]
        for loc in desired_locales:
            try:
                locale.setlocale(locale.LC_ALL, loc)
                break
            except locale.Error:
                continue
        else:
            print("Warning: Could not set German locale.")

        # Format date
        donation_date = data["date"]
        current_date = datetime.now().strftime("%d.%m.%Y")

        # Create replacements dictionary
        replacements = {
            "<<NAME>>": data["matched_name"].strip(),
            "<<STRASSE>>": data["street"].strip(),
            "<<PLZ>>": str(data["postal_code"]).strip(),
            "<<ORT>>": data["city"].strip(),
            "<<BETRAG>>": f"{float(data['amount']):.2f}".replace(".", ",") + " EUR",
            "<<BETRAG_WORTE>>": self.amount_to_words(float(data["amount"])),
            "<<DATUM_SPENDE>>": donation_date,
            "<<DATUM_HEUTE>>": current_date,
        }

        for key, value in replacements.items():
            if len(value) > 50:
                messagebox.showwarning(
                    "Warning",
                    f"Possible issue in Spendenbescheinigung_{data["matched_name"]}_{donation_date}.docx: Line break in {key} : '{value}'.",
                )

        def replace_text_in_paragraph(paragraph, replacements):
            """Helper function to replace text while preserving formatting."""
            # Store initial formatting
            runs_formatting = []
            for run in paragraph.runs:
                runs_formatting.append(
                    {
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font.name": run.font.name,
                        "font.size": run.font.size,
                        "font.color.rgb": (
                            run.font.color.rgb if run.font.color else None
                        ),
                    }
                )

            # Get the full paragraph text
            text = paragraph.text

            # Apply all replacements
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, str(value))

            # Clear the paragraph
            for run in paragraph.runs:
                run.text = ""

            # Add the new text back with original formatting
            paragraph.runs[0].text = text

            # Restore formatting
            for run, formatting in zip(paragraph.runs, runs_formatting):
                run.bold = formatting["bold"]
                run.italic = formatting["italic"]
                run.underline = formatting["underline"]
                run.font.name = formatting["font.name"]
                run.font.size = formatting["font.size"]
                if formatting["font.color.rgb"]:
                    run.font.color.rgb = formatting["font.color.rgb"]

        # Replace text in all paragraphs
        for paragraph in doc.paragraphs:
            if any(key in paragraph.text for key in replacements.keys()):
                replace_text_in_paragraph(paragraph, replacements)

        # Also check tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(key in paragraph.text for key in replacements.keys()):
                            replace_text_in_paragraph(paragraph, replacements)

        # Save document
        safe_name = "".join(x for x in data["matched_name"].strip() if x.isalnum())
        filename = f"Spendenbescheinigung_{safe_name}_{donation_date}.docx"
        doc.save(os.path.join(output_dir, filename))

        # log the receipt
        receipt_data = {
                        'generation_date': replacements['<<DATUM_HEUTE>>'],
                        'donor_name': replacements['<<NAME>>'],
                        'street': replacements['<<STRASSE>>'],
                        'postal_code': replacements['<<PLZ>>'],
                        'city': replacements['<<ORT>>'],
                        'amount': replacements['<<BETRAG>>'],
                        'amount_words': replacements['<<BETRAG_WORTE>>'],
                        'donation_date': replacements['<<DATUM_SPENDE>>'],
                        'filename': filename
                    }
        self.log_receipt(receipt_data)


    def amount_to_words(self, amount):
        """Convert amount to German words"""
        from num2words import num2words
        import math

        euros = int(amount)
        cents = int(round((amount - euros) * 100))

        if euros == 1:
            euro_str = "ein Euro"
        else:
            euro_str = num2words(euros, lang="de") + " Euro"

        if cents > 0:
            cent_str = num2words(cents, lang="de") + " Cent"
            return f"{euro_str} und {cent_str}"
        return euro_str

    def on_close(self):
        """Action to perform when closing the main window"""
        self.save_config()
        root.destroy()


class EditDialog(tk.Toplevel):
    """Dialog for editing or adding entries"""

    def __init__(self, parent, address_df, values=None):
        super().__init__(parent)
        self.title("Edit Entry" if values else "Add Entry")
        self.result = None
        
        # Make dialog resizable
        self.resizable(True, True)
        
        # Create main frame
        main_frame = ttk.Frame(self)
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Configure grid weights for resizing
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        main_frame.grid_columnconfigure(1, weight=1)

        # Create and populate fields
        fields = [
            "Name on Bank Statement",
            "Matched Name",
            "Street",
            "Postal Code",
            "City",
            "Amount (EUR)",
            "Date",
            "Match Score",
            "Purpose",
        ]
        self.entries = {}

        # Create a frame for address fields and search button
        address_frame = ttk.Frame(main_frame)
        address_frame.grid(row=1, column=0, columnspan=2, sticky="ew")
        address_frame.grid_columnconfigure(1, weight=1)

        for i, field in enumerate(fields):
            entry = ttk.Entry(main_frame)
            # ignore name on bank statement, match score and purpose
            if i > 0 and i < len(fields) - len(["Match Score", "Purpose"]):
                label = ttk.Label(main_frame, text=field, anchor="e")
                label.grid(row=i, column=0, padx=(10, 5), pady=5, sticky="e")
                entry.grid(row=i, column=1, padx=(5, 10), pady=5, sticky="ew")
                
                # Make entry wider (doubled size)
                entry.configure(width=40)  # Default was typically around 20

            if values:
                # insert bank statement name in case of not match
                if field == 'Matched Name' and values[i] == '':
                    entry.insert(0, values[i-1])
                else:
                    entry.insert(0, values[i])

            self.entries[field] = entry

        # Buttons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=len(fields), column=0, columnspan=3, pady=10)

        ttk.Button(btn_frame, text="Search Address List", command=lambda: self.search_address_list(address_df)).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Save", command=self.save).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.cancel).pack(side=tk.LEFT, padx=5)

        # Set minimum size
        self.update_idletasks()
        min_width = max(240, self.winfo_width())
        min_height = max(240, self.winfo_height())
        self.minsize(min_width, min_height)

        # Make dialog modal
        self.transient(parent)
        self.grab_set()

        # Center the dialog
        self.center_dialog(parent)

    def center_dialog(self, parent):
        """Center the dialog relative to parent window"""
        self.update_idletasks()
        parent_x = parent.winfo_rootx()
        parent_y = parent.winfo_rooty()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()

        dialog_width = self.winfo_width()
        dialog_height = self.winfo_height()

        x = parent_x + (parent_width - dialog_width) // 2
        y = parent_y + (parent_height - dialog_height) // 2

        self.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

    def search_address_list(self, address_df):
        """Open search dialog for address list"""
        # if not hasattr(parent, "address_df") or parent.address_df is None:
        if address_df is None:
            messagebox.showerror("Error", "No address data loaded!")
            return
            
        search_dialog = AddressSearchDialog(self, address_df)
        self.wait_window(search_dialog)
        
        if search_dialog.selected_address is not None:
            # Update address fields with selected data
            self.entries["Matched Name"].delete(0, tk.END)
            self.entries["Matched Name"].insert(0, search_dialog.selected_address["Name"])
            
            self.entries["Street"].delete(0, tk.END)
            self.entries["Street"].insert(0, search_dialog.selected_address["Straße"])
            
            self.entries["Postal Code"].delete(0, tk.END)
            self.entries["Postal Code"].insert(0, search_dialog.selected_address["PLZ"])
            
            self.entries["City"].delete(0, tk.END)
            self.entries["City"].insert(0, search_dialog.selected_address["Ort"])

    def save(self):
        """Save the edited values"""
        self.result = [entry.get().replace(",", ".") for entry in self.entries.values()]
        self.destroy()

    def cancel(self):
        """Cancel the edit"""
        self.destroy()

class AddressSearchDialog(tk.Toplevel):
    """Dialog for searching and selecting addresses"""
    
    def __init__(self, parent, address_df):
        super().__init__(parent)
        self.title("Search Address List")
        self.address_df = address_df
        self.selected_address = None
        
        # Make dialog resizable
        self.resizable(True, True)
        
        # Configure grid weights
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Search frame
        search_frame = ttk.Frame(self)
        search_frame.grid(row=0, column=0, padx=10, pady=5, sticky="ew")
        search_frame.grid_columnconfigure(1, weight=1)
        
        ttk.Label(search_frame, text="Search:").grid(row=0, column=0, padx=(0, 5))
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self.update_list)
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var)
        search_entry.grid(row=0, column=1, sticky="ew")
        
        # Create Treeview
        self.tree = ttk.Treeview(
            self,
            columns=("name", "street", "postal_code", "city"),
            show="headings",
            selectmode="browse"
        )
        
        # Configure columns
        self.tree.heading("name", text="Name")
        self.tree.heading("street", text="Street")
        self.tree.heading("postal_code", text="Postal Code")
        self.tree.heading("city", text="City")
        
        # Configure column widths
        self.tree.column("name", width=150)
        self.tree.column("street", width=150)
        self.tree.column("postal_code", width=100)
        self.tree.column("city", width=100)
        
        # Add scrollbars
        y_scroll = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self.tree.yview)
        x_scroll = ttk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)
        
        # Grid layout
        self.tree.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)
        y_scroll.grid(row=1, column=1, sticky="ns")
        x_scroll.grid(row=2, column=0, sticky="ew")
        
        # Buttons
        btn_frame = ttk.Frame(self)
        btn_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        ttk.Button(btn_frame, text="Select", command=self.select_address).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.cancel).pack(side=tk.LEFT, padx=5)
        
        # Initial population of list
        self.update_list()
        
        # Make dialog modal
        self.transient(parent)
        self.grab_set()
        
        # Set minimum size and center
        self.update_idletasks()
        self.minsize(600, 400)
        self.center_dialog(parent)
        
        # Set focus to search entry
        search_entry.focus_set()
        
    def center_dialog(self, parent):
        """Center the dialog relative to parent window"""
        self.update_idletasks()
        parent_x = parent.winfo_rootx()
        parent_y = parent.winfo_rooty()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()

        dialog_width = self.winfo_width()
        dialog_height = self.winfo_height()

        x = parent_x + (parent_width - dialog_width) // 2
        y = parent_y + (parent_height - dialog_height) // 2

        self.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")
        
    def update_list(self, *args):
        """Update the list based on search criteria"""
        # Clear current items
        for item in self.tree.get_children():
            self.tree.delete(item)
            
        search_term = self.search_var.get().lower()
        
        # Filter and add matching items
        for _, row in self.address_df.iterrows():
            if (search_term == "" or 
                search_term in str(row["Name"]).lower() or
                search_term in str(row["Straße"]).lower() or
                search_term in str(row["PLZ"]).lower() or
                search_term in str(row["Ort"]).lower()):
                
                self.tree.insert("", "end", values=(
                    row["Name"],
                    row["Straße"],
                    row["PLZ"],
                    row["Ort"]
                ))
    
    def select_address(self):
        """Handle address selection"""
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("Warning", "Please select an address.")
            return
            
        # Get selected item's values
        values = self.tree.item(selected_items[0])["values"]
        self.selected_address = {
            "Name": values[0],
            "Straße": values[1],
            "PLZ": values[2],
            "Ort": values[3]
        }
        self.destroy()
        
    def cancel(self):
        """Cancel selection"""
        self.selected_address = None
        self.destroy()

class ProgressDialog(tk.Toplevel):
    """Dialog showing progress during receipt generation"""

    def __init__(self, parent, max_value):
        super().__init__(parent)
        self.title("Generating Receipts")

        # Set size and position
        window_width = 500
        window_height = 100
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        center_x = int(screen_width / 2 - window_width / 2)
        center_y = int(screen_height / 2 - window_height / 2)
        self.geometry(f"{window_width}x{window_height}+{center_x}+{center_y}")

        # Configure dialog
        self.transient(parent)
        self.grab_set()

        # Add progress bar
        self.progress = ttk.Progressbar(
            self, length=300, mode="determinate", maximum=max_value
        )
        self.progress.pack(padx=20, pady=20)

        # Add label
        self.label = ttk.Label(self, text="Generating receipts...")
        self.label.pack(pady=10)

    def update(self, value):
        """Update progress bar"""
        self.progress["value"] = value
        self.label.config(text=f"Generated {value} receipts...")
        self.update_idletasks()


class LoadingProgressDialog(tk.Toplevel):
    """Dialog showing progress during data loading"""

    def __init__(self, parent):
        super().__init__(parent)
        self.title("Loading Data")

        # Set size and position
        window_width = 500
        window_height = 100
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        center_x = int(screen_width / 2 - window_width / 2)
        center_y = int(screen_height / 2 - window_height / 2)
        self.geometry(f"{window_width}x{window_height}+{center_x}+{center_y}")

        # Make dialog modal
        self.transient(parent)
        self.grab_set()

        # Prevent closing
        self.protocol("WM_DELETE_WINDOW", lambda: None)

        # Add progress bar
        self.progress = ttk.Progressbar(
            self, length=350, mode="determinate", maximum=100
        )
        self.progress.pack(padx=20, pady=(20, 10))

        # Add status label
        self.status_label = ttk.Label(self, text="Starting...")
        self.status_label.pack(pady=(0, 20))

    def update_status(self, message: str, progress_value: float):
        """Update progress bar and status message"""
        self.status_label.config(text=message)
        self.progress["value"] = progress_value
        self.update_idletasks()

class PdfConvertProgressDialog(tk.Toplevel):
    """Dialog showing progress during file processing"""
    def __init__(self, parent, max_value):
        super().__init__(parent)
        
        # Set size and position
        window_width = 400
        window_height = 100
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        center_x = int(screen_width/2 - window_width/2)
        center_y = int(screen_height/2 - window_height/2)
        self.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')
        
        # Make dialog modal
        self.transient(parent)
        self.grab_set()
        
        # Prevent closing
        self.protocol("WM_DELETE_WINDOW", lambda: None)
        
        # Add progress bar
        self.progress = ttk.Progressbar(
            self, 
            length=350, 
            mode='determinate',
            maximum=max_value
        )
        self.progress.pack(padx=20, pady=(20, 10))
        
        # Add status label
        self.label = ttk.Label(self, text="Starting conversion...")
        self.label.pack(pady=(0, 20))
        
    def update(self, value):
        """Update progress bar"""
        self.progress['value'] = value
        self.update_idletasks()

if __name__ == "__main__":
    root = tk.Tk()
    app = DonationReceiptApp(root)
    root.mainloop()
