import pandas as pd
from thefuzz import fuzz
from docx import Document
from num2words import num2words
from datetime import datetime
import locale
import argparse
import os
import math
import io
import msoffcrypto
import openpyxl
import csv
import time
# from docx2pdf import convert
from concurrent.futures import ThreadPoolExecutor
from tqdm import tqdm

def convert_to_pdf(docx_path, output_dir):
    """
    Convert a single Word document to PDF.
    
    Args:
        docx_path (str): Path to the Word document
        output_dir (str): Directory where PDF should be saved
    Returns:
        str: Path to the generated PDF file
    """
    try:
        # Create PDF filename from Word filename
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        # convert(docx_path, pdf_path)
        
        return pdf_path
    except Exception as e:
        print(f"Error converting {docx_path} to PDF: {str(e)}")
        return None

def batch_convert_to_pdf(output_dir):
    """
    Convert all Word documents in the output directory to PDF.
    
    Args:
        output_dir (str): Directory containing Word documents
    Returns:
        tuple: (successful_conversions, failed_conversions)
    """
    pdf_dir = os.path.join(output_dir, 'pdf')
    os.makedirs(pdf_dir, exist_ok=True)
    
    # Get all Word documents in the directory
    docx_files = [f for f in os.listdir(output_dir) if f.endswith('.docx')]
    
    successful = []
    failed = []
    
    print("\nConverting Word documents to PDF...")
    
    # Use ThreadPoolExecutor for parallel processing
    with ThreadPoolExecutor() as executor:
        # Create a list of futures
        futures = []
        for docx_file in docx_files:
            docx_path = os.path.join(output_dir, docx_file)
            future = executor.submit(convert_to_pdf, docx_path, pdf_dir)
            futures.append((docx_file, future))
        
        # Process results with progress bar
        for docx_file, future in tqdm(futures, desc="Converting", unit="file"):
            try:
                pdf_path = future.result()
                if pdf_path:
                    successful.append(docx_file)
                else:
                    failed.append(docx_file)
            except Exception as e:
                print(f"\nError converting {docx_file}: {str(e)}")
                failed.append(docx_file)
    
    return successful, failed

def load_and_prepare_bank_data(csv_path):
    """Load bank CSV file and extract relevant columns with proper encoding handling."""
    encodings = ['utf-8', 'iso-8859-1', 'cp1252', 'latin1']
    
    for encoding in encodings:
        try:
            with open(csv_path, 'r', encoding=encoding) as file:
                first_line = file.readline()
                delimiter = ';' if ';' in first_line else ','
            
            df = pd.read_csv(csv_path, 
                           sep=delimiter, 
                           encoding=encoding,
                           decimal=',',
                           thousands='.')
            
            print(f"Successfully read bank data with {encoding} encoding")
            return df[['Buchungstag', 'Beguenstigter/Zahlungspflichtiger', 'Betrag']]
        except UnicodeDecodeError:
            print(f"Failed to read with {encoding} encoding, trying next...")
            continue
        except Exception as e:
            print(f"Error with {encoding} encoding: {str(e)}")
            continue
    
    raise ValueError("Could not read the CSV file with any of the attempted encodings")

def load_address_data(excel_path, password=None):
    """
    Load address data from password-protected Excel file using msoffcrypto.
    
    Args:
        excel_path (str): Path to the Excel file
        password (str): Password for the protected file
    Returns:
        pandas.DataFrame: The loaded address data
    """
    try:        
        # Create a BytesIO object for the decrypted content
        decrypted_workbook = io.BytesIO()
        
        # Open and decrypt the file
        with open(excel_path, 'rb') as file:
            office_file = msoffcrypto.OfficeFile(file)
            if password:
                office_file.load_key(password=password)
            office_file.decrypt(decrypted_workbook)
        
        # Load the decrypted workbook
        workbook = openpyxl.load_workbook(filename=decrypted_workbook)
        sheet = workbook.active
        
        # Convert to pandas DataFrame
        data = []
        headers = []
        
        # Get headers from first row
        for cell in sheet[1]:
            headers.append(cell.value)
        
        # Get data from remaining rows
        for row in sheet.iter_rows(min_row=2):
            row_data = {}
            for header, cell in zip(headers, row):
                row_data[header] = cell.value
            data.append(row_data)
        
        return pd.DataFrame(data)
        
    except Exception as e:
        print(f"Error reading Excel file: {str(e)}")
        raise

def split_multiple_names(full_name):
    """
    Split a string containing multiple names into separate names.
    Handles various formats and separators.
    """
    full_name = str(full_name).strip()
    
    # List of possible separators
    separators = [' Und ', ' und ', ' U. ', ' u. ', ' And ', ' and ', ' & ', ' + ', '   ']
    
    # First try explicit separators
    for sep in separators:
        if sep in full_name:
            split_names = full_name.split(sep)
            if len(split_names[0].split(' ')) == 1:
                split_names = [split_names[0] + ' ' + split_names[-1].split(' ')[-1], split_names[1]]
            return [name.strip() for name in split_names]
    
    # If no explicit separator, try to detect multiple full names
    # by looking for patterns like multiple last names
    words = full_name.split(' ')
    if len(words) >= 4:  # Minimum 4 words needed for 2 full names
        # Try to find repeated last names
        last_name = words[-1]
        for i in range(len(words)-2, 0, -1):
            if words[i] == last_name:
                return [
                    ' '.join(words[:i]),
                    ' '.join(words[i:])
                ]
    
    # If no pattern found, return as single name
    return [full_name]

def normalize_name(name):
    """
    Normalize name for comparison, handling different formats.
    """
    name = str(name).strip()
    
    # Handle "last_name, first_name" format
    if ',' in name:
        parts = name.split(',')
        if len(parts) == 2:
            last_name = parts[0].strip()
            first_name = parts[1].strip()
            return f"{first_name} {last_name}"

    return name

def find_best_match(donor_name, address_df, threshold=80):
    """
    Find the best matching address using fuzzy matching.
    Handles multiple names and tries various matching strategies.
    """
    best_score = 0
    best_match = None
    has_best_match = False
    original_donor_name = donor_name
    matched_name = None

    # turn posible all caps into regular title format
    formatted_name = donor_name.title()
    # normalize "last_name, firstname" format
    donor_name_normalized = normalize_name(formatted_name)
    
    # Split into potential multiple names
    donor_names = split_multiple_names(donor_name_normalized)
    print('')
    if len(donor_names) > 1:
        print(f"Split '{donor_name_normalized}' into: {donor_names}")
    
    # Try matching each name individually and combined
    for name in donor_names:
        normalized_name = normalize_name(name)
        print(normalized_name)
        
        for _, row in address_df.iterrows():
            list_name_raw = str(row['Name'])
            list_name_normalized_raw = normalize_name(list_name_raw)

            list_names = split_multiple_names(list_name_raw)
            list_names_normalized = split_multiple_names(list_name_normalized_raw)

            for [list_name, list_name_normalized] in zip(list_names, list_names_normalized):
                # Try different matching combinations
                scores = [
                    fuzz.ratio(name.lower(), list_name.lower()),
                    fuzz.ratio(normalized_name.lower(), list_name.lower()),
                    fuzz.ratio(normalized_name.lower(), list_name_normalized.lower()),
                    fuzz.token_sort_ratio(name.lower(), list_name.lower()),
                    fuzz.token_sort_ratio(normalized_name.lower(), list_name_normalized.lower())
                ]
                
                max_score = max(scores)
            
                if max_score > best_score and max_score >= threshold:
                    best_score = max_score
                    best_match = row
                    has_best_match = True
                    matched_name = name

    if len(donor_names) > 1 and not has_best_match:
        # If no match found and we have multiple names, 
        # try matching the combined names
        combined_name = ' '.join(donor_names)
        
        for _, row in address_df.iterrows():
            list_name = str(row['Name'])
            list_name_normalized = normalize_name(list_name)

            scores = [
                fuzz.ratio(combined_name.lower(), list_name.lower()),
                fuzz.ratio(combined_name.lower(), list_name_normalized.lower()),
                fuzz.token_sort_ratio(combined_name.lower(), list_name.lower()),
                fuzz.token_sort_ratio(combined_name.lower(), list_name_normalized.lower())
            ]
            
            max_score = max(scores)
            
            if max_score > best_score and max_score >= threshold:
                best_score = max_score
                best_match = row
                matched_name = combined_name
    
    if best_match is not None:
        print(f"Found match for '{matched_name}' in '{best_match['Name']}' with score {best_score}")
        if matched_name != original_donor_name:
            print(f"Note: Matched using partial name from '{original_donor_name}'")
    else:
        print(f"No match found for any name in '{original_donor_name}'")
    
    return best_match, best_score

def format_date(date_int):
    """
    Convert date from integer format (DDMMYY) to German date string (DD.MM.YYYY).
    
    Args:
        date_int (int): Date as integer in format DDMMYY (e.g., 130125 for 13.01.25)
    Returns:
        str: Date in format 'DD.MM.YYYY'
    """
    try:
        # Convert integer to string and pad with leading zeros if necessary
        date_str = str(date_int)
        if len(date_str) < 6:
          date_str = str(date_int).zfill(6)
        elif len(date_str) == 7:
          date_str = str(date_int).zfill(8)

        # Extract components
        day = date_str[:2]
        month = date_str[2:4]
        year = date_str[4:]
        
        # Convert two-digit year to four-digit year
        year_int = int(year)
        if year_int < 100:
            year = f"20{year}"
        
        # Return formatted date
        return f"{day}.{month}.{year}"
        
    except Exception as e:
        print(f"Error formatting date {date_int}: {str(e)}")
        return str(date_int)

def amount_to_words(amount) -> str:
    cents, euros = math.modf(amount)

    if round(euros) == 1:
        amount_str = 'ein Euro'    
    else:
        amount_str = num2words(euros, lang='de') + ' Euro'

    if cents > 0:
        amount_str += ' und ' + num2words(round(cents * 100), lang='de') + ' Cent'

    return amount_str

def generate_receipt(template_path, donor_info, amount, transaction_date):
    """Generate donation receipt from template."""
    try:
        doc = Document(template_path)

        desired_locales = ['de_DE.UTF-8', 'de_DE', 'de_de', 'German']
        for loc in desired_locales:
            try:
                locale.setlocale(locale.LC_ALL, loc)
                break
            except locale.Error:
                continue
        else:
            print("Warning: Could not set German locale.")
        
        replacements = {
            '<<NAME>>': donor_info['Name'].strip(),
            '<<STRASSE>>': donor_info['Straße'].strip(),
            '<<PLZ>>': str(donor_info['PLZ']).strip(),
            '<<ORT>>': donor_info['Ort'].strip(),
            '<<BETRAG>>': f'{amount:.2f}'.replace('.', ',') + ' EUR',
            '<<BETRAG_WORTE>>': amount_to_words(amount),
            '<<DATUM_SPENDE>>': format_date(transaction_date),
            '<<DATUM_HEUTE>>': datetime.now().strftime('%d.%m.%Y')
        }

        for key, value in replacements.items():
            if len(value) > 50:
                print(f" ! WARNING ! Possible line break for {key}: {value}")
        
        def replace_text_in_paragraph(paragraph, replacements):
            """Helper function to replace text while preserving formatting."""
            # Store initial formatting
            runs_formatting = []
            for run in paragraph.runs:
                runs_formatting.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font.name': run.font.name,
                    'font.size': run.font.size,
                    'font.color.rgb': run.font.color.rgb if run.font.color else None
                })
            
            # Get the full paragraph text
            text = paragraph.text
            
            # Apply all replacements
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, str(value))
            
            # Clear the paragraph
            for run in paragraph.runs:
                run.text = ""
            
            # Add the new text back with original formatting
            paragraph.runs[0].text = text
            
            # Restore formatting
            for run, formatting in zip(paragraph.runs, runs_formatting):
                run.bold = formatting['bold']
                run.italic = formatting['italic']
                run.underline = formatting['underline']
                run.font.name = formatting['font.name']
                run.font.size = formatting['font.size']
                if formatting['font.color.rgb']:
                    run.font.color.rgb = formatting['font.color.rgb']
        
        # Replace text in all paragraphs
        for paragraph in doc.paragraphs:
            if any(key in paragraph.text for key in replacements.keys()):
                replace_text_in_paragraph(paragraph, replacements)
        
        # Also check tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(key in paragraph.text for key in replacements.keys()):
                            replace_text_in_paragraph(paragraph, replacements)
        
        return doc
    except Exception as e:
        print(f"Error generating receipt: {str(e)}")
        raise

def create_receipt_log(file_name, output_dir):
    """Create a CSV file for logging receipt information."""
    log_file = os.path.join(output_dir, file_name)
    headers = [
        'Date Generated',
        'Donor Name',
        'Street',
        'Postal Code',
        'City',
        'Donation Amount',
        'Donation Amount Words',
        'Donation Date',
        'Match Score',
        'Receipt Filename'
    ]
    
    with open(log_file, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.writer(f, delimiter=';')
        writer.writerow(headers)
    
    return log_file

def log_receipt(log_file, receipt_data):
    """Log receipt information to CSV file."""
    with open(log_file, 'a', newline='', encoding='utf-8-sig') as f:
        writer = csv.writer(f, delimiter=';')
        writer.writerow([
            receipt_data['generation_date'],
            receipt_data['donor_name'],
            receipt_data['street'],
            receipt_data['postal_code'],
            receipt_data['city'],
            receipt_data['amount'],
            receipt_data['amount_words'],
            receipt_data['donation_date'],
            receipt_data['match_score'],
            receipt_data['filename']
        ])

def process_donations(args):
    """Main function to process all donations and generate receipts."""
    try:
        # Create output directory if it doesn't exist
        os.makedirs(args.output_dir, exist_ok=True)
        
        # Create receipt log file
        log_file = create_receipt_log(args.output_log, args.output_dir)
        
        # Load data
        print("Loading bank data...")
        bank_data = load_and_prepare_bank_data(args.bank_csv)
        print("Loading address data...")
        address_data = load_address_data(args.address_excel, password=args.password)
        
        # Process each donation
        total_processed = 0
        total_matched = 0
        no_matches = []
        
        print("\nProcessing donations...")
        for _, donation in bank_data.iterrows():
            try:
                if donation['Betrag'] <= 0:
                    continue

                donor_name = donation['Beguenstigter/Zahlungspflichtiger']
                amount_str = str(donation['Betrag']).replace(',', '.')
                amount = float(amount_str)
                transaction_date = donation['Buchungstag']

                # Find matching address
                donor_info, match_score = find_best_match(donor_name, address_data, args.threshold)
                
                if donor_info is not None:
                    # Generate receipt
                    receipt = generate_receipt(args.template, donor_info, amount, transaction_date)
                    
                    # Generate filename
                    safe_name = "".join(x for x in donor_info['Name'].strip() if x.isalnum())
                    filename = f'Spendenbescheinigung_{safe_name}_{format_date(transaction_date)}.docx'
                    full_path = os.path.join(args.output_dir, filename)
                    
                    # Save receipt
                    receipt.save(full_path)
                    
                    # Prepare receipt data for logging
                    receipt_data = {
                        'generation_date': datetime.now().strftime('%d.%m.%Y'),
                        'donor_name': donor_info['Name'].strip(),
                        'street': donor_info['Straße'].strip(),
                        'postal_code': str(donor_info['PLZ']).strip(),
                        'city': donor_info['Ort'].strip(),
                        'amount': f'{amount:.2f}'.replace('.', ',') + ' EUR',
                        'amount_words': amount_to_words(amount),
                        'donation_date': format_date(transaction_date),
                        'match_score': match_score,
                        'filename': filename
                    }
                    
                    # Log receipt information
                    log_receipt(log_file, receipt_data)
                    
                    print(f'Generated receipt for {donor_name} (match score: {match_score}%)')
                    total_matched += 1
                else:
                    no_matches.append(donor_name)
                
                total_processed += 1
            
            except Exception as e:
                print(f"Error processing donation for {donor_name}: {str(e)}")
                continue
        
        # TODO: Convert all generated Word documents to PDF; this needs docx2pdf running on WSL
        # successful_conversions, failed_conversions = batch_convert_to_pdf(args.output_dir)

        # Print summary
        print(f"\nProcessing complete!")
        print(f"Total donations processed: {total_processed}")
        print(f"Successfully matched and generated: {total_matched}")
        print(f"Could not find matches for: {len(no_matches)} donations")
        print(f"\nReceipt log saved to: {log_file}")
        
        if no_matches:
            print("\nDonors with no matching address found:")
            for name in no_matches:
                print(f"- {name}")
    
    except Exception as e:
        print(f"Error in main process: {str(e)}")
        raise

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate donation receipts from bank data and address list')
    
    parser.add_argument('--bank-csv', required=True,
                      help='Path to bank CSV file')
    parser.add_argument('--address-excel', required=True,
                      help='Path to address Excel file')
    parser.add_argument('--template', required=True,
                      help='Path to Word template file')
    parser.add_argument('--output-dir', required=True,
                      help='Output directory for generated receipts')
    parser.add_argument('--output-log',
                      help='CSV file to log the generated receips to.')
    parser.add_argument('--password',
                      help='Password for protected Excel file', 
                      default=None)
    parser.add_argument('--threshold', type=int,
                      help='Matching threshold (0-100)', 
                      default=80)
    
    args = parser.parse_args()
    
    process_donations(args)
